"""
extractor.py
------------
Pulls raw text out of PDF, DOCX, and TXT files.

Supports two input modes:
    1. File path (str)    — file already on disk
    2. Raw bytes          — from FastAPI UploadFile, no temp file needed

Change based on SBERT teammate feedback:
    - DOCX table extraction now skips header-only rows (single row tables
      or first rows that look like column labels e.g. "Factor  Effect on Ecosystem")
      These get caught again in cleaner.py but filtering early is cleaner.
"""

from __future__ import annotations

import io
import os
from typing import Tuple, Dict, Any, Optional, Union

import fitz   # PyMuPDF
import docx   # python-docx

from .exceptions import UnsupportedFileTypeError, ExtractionError
from .validators import validate_filename

_TXT_ENCODING_FALLBACKS = ("utf-8-sig", "utf-8", "cp1252", "latin-1")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _extract_pdf_pages(doc: "fitz.Document", label: str) -> Tuple[str, Dict[str, Any]]:
    page_texts = []
    per_page_char_counts = []

    for page in doc:
        text = page.get_text("text")
        page_texts.append(text)
        per_page_char_counts.append(len(text))

    if doc.page_count == 0:
        raise ExtractionError(f"PDF '{label}' has no pages.")

    full_text = "\n".join(page_texts)

    if not full_text.strip():
        raise ExtractionError(
            f"PDF '{label}' produced no extractable text. "
            "It may be a scanned/image-only document requiring OCR."
        )

    return full_text, {
        "file_type": "pdf",
        "num_pages": doc.page_count,
        "per_page_char_counts": per_page_char_counts,
    }


def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"PDF not found: {file_path}")
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise ExtractionError(f"Could not open PDF '{file_path}': {e}") from e
    try:
        if doc.is_encrypted and not doc.authenticate(""):
            raise ExtractionError(f"PDF '{file_path}' is password-protected.")
        return _extract_pdf_pages(doc, file_path)
    finally:
        doc.close()


def extract_text_from_pdf_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        raise ExtractionError(f"Could not open in-memory PDF: {e}") from e
    try:
        if doc.is_encrypted and not doc.authenticate(""):
            raise ExtractionError("In-memory PDF is password-protected.")
        return _extract_pdf_pages(doc, "<in-memory PDF>")
    finally:
        doc.close()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _is_table_header_row(cells: list) -> bool:
    """
    Heuristic: a table row is a header row if it has short cells with
    no sentence-ending punctuation -- e.g. ["Factor", "Effect on Ecosystem"].
    We skip these rows since they are structural labels, not content.
    """
    if not cells:
        return False
    texts = [c.strip() for c in cells if c.strip()]
    if not texts:
        return False
    # All cells are short and none end with sentence punctuation
    return all(len(t.split()) <= 6 and not t[-1] in ".?!" for t in texts)


def _extract_docx_content(document: "docx.Document", label: str) -> Tuple[str, Dict[str, Any]]:
    text_chunks = []

    # Paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            text_chunks.append(para.text)

    # Tables — skip header rows, only extract content rows
    num_tables = len(document.tables)
    for table in document.tables:
        for row_idx, row in enumerate(table.rows):
            cell_texts = [cell.text.strip() for cell in row.cells]

            # Skip the first row if it looks like a header
            if row_idx == 0 and _is_table_header_row(cell_texts):
                continue

            for cell_text in cell_texts:
                if cell_text:
                    text_chunks.append(cell_text)

    full_text = "\n".join(text_chunks)
    if not full_text.strip():
        raise ExtractionError(f"DOCX '{label}' has no extractable text.")

    return full_text, {
        "file_type": "docx",
        "num_paragraphs": len(document.paragraphs),
        "num_tables": num_tables,
    }


def extract_text_from_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"DOCX not found: {file_path}")
    try:
        document = docx.Document(file_path)
    except Exception as e:
        raise ExtractionError(f"Could not open DOCX '{file_path}': {e}") from e
    return _extract_docx_content(document, file_path)


def extract_text_from_docx_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not open in-memory DOCX: {e}") from e
    return _extract_docx_content(document, "<in-memory DOCX>")


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------

def _decode_txt_bytes(file_bytes: bytes, label: str) -> Tuple[str, Dict[str, Any]]:
    encodings = list(_TXT_ENCODING_FALLBACKS)
    if not file_bytes.startswith(b"\xef\xbb\xbf"):
        encodings = [e for e in encodings if e != "utf-8-sig"]

    text = None
    used_encoding = None
    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            used_encoding = encoding
            break
        except (UnicodeDecodeError, LookupError):
            continue

    if text is None:
        raise ExtractionError(
            f"Could not decode TXT '{label}' with any supported encoding."
        )
    if not text.strip():
        raise ExtractionError(f"TXT '{label}' is empty.")

    return text, {
        "file_type": "txt",
        "encoding_used": used_encoding,
        "num_lines": text.count("\n") + 1,
    }


def extract_text_from_txt(file_path: str) -> Tuple[str, Dict[str, Any]]:
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"TXT not found: {file_path}")
    with open(file_path, "rb") as f:
        return _decode_txt_bytes(f.read(), file_path)


def extract_text_from_txt_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    return _decode_txt_bytes(file_bytes, "<in-memory TXT>")


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def extract_text(
    source: Union[str, bytes],
    filename: Optional[str] = None,
) -> Tuple[str, Dict[str, Any]]:
    """
    Single entry point — dispatches to the right extractor based on extension.

    Args:
        source: file path (str) OR raw bytes
        filename: required when source is bytes (e.g. "essay.pdf")

    Raises:
        UnsupportedFileTypeError, ValueError, FileNotFoundError, ExtractionError
    """
    if isinstance(source, (bytes, bytearray)):
        if not filename:
            raise ValueError(
                "filename is required when passing raw bytes "
                "(e.g. filename='essay.pdf') so the file type can be determined."
            )
        validate_filename(filename)
        ext = os.path.splitext(filename)[1].lower()
        dispatch_bytes = {
            ".pdf":  extract_text_from_pdf_bytes,
            ".docx": extract_text_from_docx_bytes,
            ".txt":  extract_text_from_txt_bytes,
        }
        return dispatch_bytes[ext](bytes(source))

    validate_filename(source)
    ext = os.path.splitext(source)[1].lower()
    dispatch_path = {
        ".pdf":  extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".txt":  extract_text_from_txt,
    }
    return dispatch_path[ext](source)
